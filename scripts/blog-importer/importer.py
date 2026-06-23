import os
import sys
import zipfile
import re
from docx import Document

def clean_slug(title):
    # Convert title to lowercase and replace non-alphanumeric chars with hyphens
    slug = title.lower().strip()
    slug = re.sub(r'[^a-z0-9\s-]', '', slug)
    slug = re.sub(r'[\s-]+', '-', slug)
    return slug

def extract_docx(docx_path, project_root):
    if not os.path.exists(docx_path):
        print(f"Error: File {docx_path} does not exist.")
        sys.exit(1)

    print(f"Reading document: {docx_path}")
    
    # Load document
    doc = Document(docx_path)
    
    # Try to find title for slug
    title = "untitled-post"
    for para in doc.paragraphs:
        if para.text.strip():
            title = para.text.strip()
            break
            
    slug = clean_slug(title)
    print(f"Generated slug: {slug}")
    
    # Setup directories
    image_dir = os.path.join(project_root, "public", "images", "blog", slug)
    os.makedirs(image_dir, exist_ok=True)
    
    # 1. Extract Images from ZIP structure
    extracted_images = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            image_idx = 1
            for file_info in z.infolist():
                if file_info.filename.startswith('word/media/'):
                    ext = os.path.splitext(file_info.filename)[1]
                    # Give them sequential clean names
                    new_name = f"image-{image_idx}{ext}"
                    out_path = os.path.join(image_dir, new_name)
                    
                    with open(out_path, 'wb') as out_f:
                        out_f.write(z.read(file_info.filename))
                    
                    # Web path to reference in Markdown
                    web_path = f"/images/blog/{slug}/{new_name}"
                    extracted_images.append((file_info.filename, web_path))
                    print(f"Extracted image to: {web_path}")
                    image_idx += 1
    except Exception as e:
        print(f"Warning failed to extract images via zip: {e}")
        
    # 2. Extract Text
    text_content = []
    text_content.append(f"=== METADATA ===")
    text_content.append(f"TITLE: {title}")
    text_content.append(f"SLUG: {slug}")
    text_content.append(f"=== IMAGES EXTRACTED ===")
    for orig, web in extracted_images:
        text_content.append(f"{orig} -> {web}")
    text_content.append(f"=== BODY TEXT ===")
    
    for para in doc.paragraphs:
        text_content.append(para.text)
        
    # Save raw text output for AI to read
    raw_text_path = os.path.join(project_root, "scripts", "blog-importer", "temp_raw_text.txt")
    with open(raw_text_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(text_content))
        
    print(f"Raw text and structure saved to: {raw_text_path}")
    print("SUCCESS")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python importer.py <path_to_docx> <project_root>")
        sys.exit(1)
        
    docx_path = sys.argv[1]
    project_root = sys.argv[2]
    extract_docx(docx_path, project_root)
